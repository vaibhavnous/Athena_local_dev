
from __future__ import annotations

import hashlib
import json
import os
import io
import traceback
import uuid
from datetime import datetime, timezone
from contextlib import redirect_stderr, redirect_stdout
from typing import Optional

import docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pinecone import Pinecone
from pydantic import ValidationError

from schema import BRDSchema
from state import Stage01State
from utilis.db import config, get_pipeline_connection
from utilis.env import load_backend_env
from utilis.logger import logger
from utilis.db import execute_source_sql

load_backend_env()
DEV_MODE = os.getenv("DEV_MODE", "").strip().lower() in {"1", "true", "yes", "on"}


pinecone_api_key = os.environ.get("PINECONE_API_KEY")
pinecone_index_name = "ai-store-index"

pc = None
pinecone_index = None

try:
    if not pinecone_api_key:
        logger.warning("Pinecone API key not found", extra={"node": "ingestion_bootstrap"})
    else:
        logger.info("Initializing Pinecone client", extra={"node": "ingestion_bootstrap"})
        pc = Pinecone(api_key=pinecone_api_key)
        if DEV_MODE:
            indexes = pc.list_indexes()
            logger.info("Available Pinecone indexes: %s", indexes, extra={"node": "ingestion_bootstrap"})
        pinecone_index = pc.Index(name=pinecone_index_name)
        logger.info("Connected to Pinecone index %s", pinecone_index_name, extra={"node": "ingestion_bootstrap"})
except Exception as e:
    logger.warning("Pinecone init failed: %s", e, extra={"node": "ingestion_bootstrap"})


try:
    # Do not initialize embeddings at import-time. Some environments have restricted
    # outbound network (HF downloads) and we still want the backend to import cleanly.
    _embedding_model = None
except Exception as e:
    # Keep a conservative fallback. Errors here should not block module import.
    logger.warning("Embedding model bootstrap skipped: %s", e, extra={"node": "ingestion_bootstrap"})
    _embedding_model = None


def _get_embedding_model(*, log_context: dict) -> Optional[HuggingFaceEmbeddings]:
    global _embedding_model
    if _embedding_model is not None:
        return _embedding_model

    try:
        if os.getenv("ATHENA_ENABLE_EMBEDDINGS", "").strip().lower() not in {"1", "true", "yes", "on"}:
            logger.info("Semantic indexing deferred; continuing with catalog-driven analysis", extra=log_context)
            return None

        from langchain_huggingface import HuggingFaceEmbeddings

        logger.info("Initializing local embedding model", extra={"node": "ingestion_bootstrap"})
        os.environ["TRANSFORMERS_NO_ADVISE"] = "1"
        os.environ["TOKENIZERS_PARALLELISM"] = "false"

        if DEV_MODE:
            _embedding_model = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={"local_files_only": True, "trust_remote_code": False},
                encode_kwargs={"normalize_embeddings": False},
            )
            _embedding_model.embed_query("hello world")
        else:
            with redirect_stdout(io.StringIO()), redirect_stderr(io.StringIO()):
                _embedding_model = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={"local_files_only": True, "trust_remote_code": False},
                    encode_kwargs={"normalize_embeddings": False},
                )
                _embedding_model.embed_query("hello world")

        return _embedding_model
    except Exception as exc:
        logger.info("Semantic indexing deferred; continuing with catalog-driven analysis: %s", exc, extra=log_context)
        _embedding_model = None
        return None


db_conf = config["azure_sql"]
pinecone_conf = config.get("pinecone", {})
pipeline_schema = db_conf.get("pipeline_schema") or db_conf.get("schema_name") or "metadata"
TOKEN_BUDGET = 50000


def _copy_state(state: Stage01State) -> Stage01State:
    return state.copy()


def _run_id(state: Stage01State, default: str = "unknown") -> str:
    return state.get("run_id") or default


def _context(state: Stage01State, node: str, run_id_default: str = "unknown") -> dict:
    return {"run_id": _run_id(state, run_id_default), "node": node}


def _mark_failed(state: Stage01State, error: str) -> Stage01State:
    state.update({
        "error": error,
        "status": "FAILED",
    })
    return state


def _mark_embedding_skipped(
    state: Stage01State,
    *,
    brd_embedded: Optional[bool] = None,
    schema_embedded: Optional[bool] = None,
    schema_columns_count: Optional[int] = None,
) -> Stage01State:
    updates = {}
    if brd_embedded is not None:
        updates["brd_embedded"] = brd_embedded
    if schema_embedded is not None:
        updates["schema_embedded"] = schema_embedded
    if schema_columns_count is not None:
        updates["schema_columns_count"] = schema_columns_count
    state.update(updates)
    return state


def _build_context_text_from_data(data: object) -> str:
    columns = getattr(data, "columns", None)
    if columns is None:
        return ""

    column_names = [str(column).strip() for column in columns if str(column).strip()]
    if not column_names:
        return ""

    return (
        "Dataset contains columns: "
        + ", ".join(column_names)
        + ". This structured dataset should be analyzed to infer business requirements, "
        + "reporting intent, and candidate KPIs from the available schema."
    )


def _parse_input(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    log_context = _context(new_state, "parse_input", "pending_generation")
    logger.info("START: _parse_input", extra=log_context)

    try:
        if state.get("data") is not None:
            context_text = _build_context_text_from_data(state.get("data"))
            if not context_text:
                logger.warning("Structured input did not expose usable columns", extra=log_context)
                return new_state

            log_context["context_length"] = len(context_text)
            logger.info("Built context_text from structured data", extra=log_context)
            new_state["context_text"] = context_text
            # Keep legacy helpers working without changing graph structure.
            new_state["brd_text"] = context_text
            logger.info("END: _parse_input", extra=log_context)
            return new_state

        input_value = state.get("brd_text", "").strip()

        if not input_value:
            logger.warning("Empty input received in parser helper", extra=log_context)
            return new_state

        if os.path.exists(input_value):
            log_context["file_path"] = input_value
            logger.info("Detected file input", extra=log_context)

            ext = os.path.splitext(input_value)[1].lower()
            if ext == ".txt":
                with open(input_value, "r", encoding="utf-8") as f:
                    parsed_text = f.read().strip()
            elif ext == ".docx":
                doc = docx.Document(input_value)
                full_text = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())

                for table in doc.tables:
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_data:
                            full_text.append(" | ".join(row_data))

                parsed_text = "\n".join(full_text)
            else:
                logger.warning("Unsupported file type: %s", ext, extra=log_context)
                parsed_text = input_value
        else:
            logger.info("Detected raw text input", extra=log_context)
            parsed_text = input_value

        log_context["parsed_length"] = len(parsed_text)
        logger.info("Successfully parsed text", extra=log_context)
        new_state["context_text"] = parsed_text
        new_state["brd_text"] = parsed_text

        logger.info("END: _parse_input", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in _parse_input", extra=log_context)
        raise


def _acquire_and_validate_brd(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    if not new_state.get("run_id"):
        new_state["run_id"] = str(uuid.uuid4())

    log_context = _context(new_state, "acquire_and_validate")
    logger.info("START: _acquire_and_validate_brd", extra=log_context)

    try:
        context_text = (new_state.get("context_text") or new_state.get("brd_text") or "").strip()
        log_context["context_length"] = len(context_text)
        logger.info("Checking input context length", extra=log_context)

        if not context_text:
            logger.error("Validation failed: input context is empty", extra=log_context)
            return _mark_failed(new_state, "Validation Failed: input context is empty.")

        if new_state.get("data") is None and len(context_text) < 200:
            warning = "BRD is shorter than the preferred 200-character minimum; continuing with a warning."
            logger.warning(warning, extra=log_context)
            new_state["validation_warning"] = warning

        new_state.update({
            "context_text": context_text,
            "brd_text": context_text,
            "status": "IN_PROGRESS",
            "error": None,
        })

        logger.info("END: _acquire_and_validate_brd", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.exception("ERROR in _acquire_and_validate_brd", extra=log_context)
        raise


def _estimate_and_fingerprint(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    log_context = _context(new_state, "estimate_and_fingerprint")
    logger.info("START: _estimate_and_fingerprint", extra=log_context)

    try:
        if state.get("status") == "FAILED":
            logger.warning("Skipping helper because status is FAILED", extra=log_context)
            return new_state

        brd_text = state.get("brd_text", "")
        token_estimate = max(1, len(brd_text) // 4)
        log_context["token_estimate"] = token_estimate
        logger.info("Token estimate calculated", extra=log_context)

        fingerprint = hashlib.sha256(brd_text.encode("utf-8")).hexdigest()
        log_context["fingerprint"] = fingerprint
        logger.info("Fingerprint generated", extra=log_context)

        new_state.update({
            "token_estimate": token_estimate,
            "fingerprint": fingerprint,
            "status": "IN_PROGRESS",
        })

        logger.info("END: _estimate_and_fingerprint", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in _estimate_and_fingerprint", extra=log_context)
        raise


def _validate_budget(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    log_context = _context(new_state, "validate_budget")
    logger.info("START: _validate_budget", extra=log_context)

    if new_state.get("status") == "FAILED":
        logger.warning("Budget helper skipping because status is FAILED", extra=log_context)
        return new_state

    token_count = new_state.get("token_estimate", 0)
    log_context["token_count"] = token_count

    if token_count > TOKEN_BUDGET:
        logger.error("Token budget exceeded (%s > %s)", token_count, TOKEN_BUDGET, extra=log_context)
        return _mark_failed(new_state, f"Token budget exceeded ({token_count} > {TOKEN_BUDGET}).")

    logger.info("END: _validate_budget", extra=log_context)
    return new_state


def _validate_pricing_config(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    log_context = _context(new_state, "validate_pricing")
    logger.info("START: _validate_pricing_config", extra=log_context)

    try:
        if new_state.get("status") == "FAILED":
            logger.warning("Skipping pricing helper due to FAILED status", extra=log_context)
            return new_state

        pricing_config = {"input_cost_per_1k": 0.01, "output_cost_per_1k": 0.03}
        token_estimate = new_state.get("token_estimate", 0)
        estimated_cost = (token_estimate / 1000) * pricing_config["input_cost_per_1k"]

        log_context["estimated_cost"] = estimated_cost
        logger.info("Estimated cost calculated", extra=log_context)

        metadata = new_state.get("metadata") or {}
        metadata["estimated_cost"] = estimated_cost
        metadata["pricing_config"] = pricing_config

        new_state.update({
            "metadata": metadata,
            "pricing_validated": True,
        })

        logger.info("END: _validate_pricing_config", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in _validate_pricing_config", extra=log_context)
        raise


def _validate_schema(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    log_context = _context(new_state, "validate_schema")
    logger.info("START: _validate_schema", extra=log_context)

    try:
        if new_state.get("status") == "FAILED":
            logger.warning("Skipping schema helper due to FAILED status", extra=log_context)
            return new_state

        try:
            if new_state.get("data") is not None:
                if not (new_state.get("context_text") or "").strip():
                    raise ValueError("Structured data context is empty")
                new_state["is_schema_valid"] = True
                logger.info("Structured-data context validation successful", extra=log_context)
            else:
                BRDSchema(content=new_state.get("context_text") or new_state.get("brd_text", ""))
                new_state["is_schema_valid"] = True
                logger.info("Schema validation successful", extra=log_context)
        except (ValidationError, ValueError) as e:
            log_context["error"] = str(e)
            logger.error("Schema validation failed", extra=log_context)
            new_state.update({
                "error": f"Schema Validation Failed: {str(e)}",
                "status": "FAILED",
                "is_schema_valid": False,
            })
            return new_state

        logger.info("END: _validate_schema", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in _validate_schema", extra=log_context)
        raise


def _store_and_register(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    run_id = _run_id(new_state, str(uuid.uuid4()))
    new_state["run_id"] = run_id

    log_context = _context(new_state, "store_and_register")
    logger.info("START: _store_and_register", extra=log_context)

    if new_state.get("status") == "FAILED":
        logger.warning("Skipping storage helper due to FAILED status", extra=log_context)
        return new_state

    try:
        fingerprint = new_state.get("fingerprint") or ""
        token_count = new_state.get("token_estimate", 0)
        metadata = new_state.get("metadata", {})
        brd_text = new_state.get("brd_text") or ""
        pipeline_status = "STAGE_01_COMPLETE"
        utc_now = datetime.now(timezone.utc)
        metadata_str = json.dumps(metadata)
        conn = get_pipeline_connection()
        try:
            cursor = conn.cursor()

            cursor.execute(
                f"""
                INSERT INTO [{pipeline_schema}].brd_run_registry
                (
                    run_id,
                    status,
                    token_count,
                    timestamp
                )
                VALUES (?, ?, ?, ?)
                """,
                (
                    run_id,
                    pipeline_status,
                    token_count,
                    utc_now,
                ),
            )

            try:
                cursor.execute(
                    f"""
                    INSERT INTO [{pipeline_schema}].pipeline_run_log
                    (run_id, source, status, started_at, completed_at, error_message)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (
                        run_id,
                        new_state.get("source", "unknown"),
                        pipeline_status,
                        utc_now,
                        utc_now,
                        None,
                    ),
                )
            except Exception as exc:
                logger.warning("Pipeline run log persistence skipped: %s", exc, extra=log_context)

            conn.commit()
        finally:
            conn.close()

        new_state.update({
            "run_id": run_id,
            "status": pipeline_status,
        })

        logger.info("Run successfully registered in Azure SQL", extra=log_context)
        logger.info("END: _store_and_register", extra=log_context)
        return new_state
    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in _store_and_register", extra=log_context)
        raise


def finalize_ingestion_after_memory(state: Stage01State) -> Stage01State:
    new_state = _validate_pricing_config(state)
    if new_state.get("status") == "FAILED":
        return new_state

    new_state = _validate_schema(new_state)
    if new_state.get("status") == "FAILED":
        return new_state

    return _store_and_register(new_state)


def _chunk_and_embed(state: Stage01State) -> Stage01State:
    new_state = state.copy()
    run_id = new_state.get("run_id", "unknown")

    log_context = {
        "run_id": run_id,
        "node": "chunk_and_embed",
    }

    logger.info("START: _chunk_and_embed", extra=log_context)

    try:
        if new_state.get("status") == "FAILED":
            logger.warning("Skipping embedding due to FAILED status", extra=log_context)
            return new_state

        model = _get_embedding_model(log_context=log_context)
        if model is None:
            logger.info("BRD semantic index deferred; requirement extraction will use parsed BRD content", extra=log_context)
            return _mark_embedding_skipped(new_state, brd_embedded=False)

        if pinecone_index is None:
            raise Exception("Pinecone not initialized")

        brd_text = new_state.get("brd_text", "").strip()
        fingerprint = new_state.get("fingerprint", "")

        if not brd_text:
            raise Exception("Empty BRD text")

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""],
        )

        docs = splitter.create_documents(
            texts=[brd_text],
            metadatas=[{
                "run_id": run_id,
                "fingerprint": fingerprint,
                "source": "BRD",
            }],
        )

        log_context["chunk_count"] = len(docs)
        logger.info("Split into %d chunks", len(docs), extra=log_context)

        if not docs:
            raise Exception("No chunks created")

        texts = [doc.page_content for doc in docs]
        logger.info("Generating embeddings...", extra=log_context)
        vectors = model.embed_documents(texts)

        index_name = pinecone_conf.get("index_name", pinecone_index_name)
        namespace = "global"

        logger.info(
            "Upserting %d chunks to Pinecone (%s, namespace=%s)",
            len(docs),
            index_name,
            namespace,
            extra=log_context,
        )

        pc = Pinecone(api_key=pinecone_conf.get("api_key") or os.getenv("PINECONE_API_KEY"))
        index = pc.Index(index_name)

        try:
            index.delete(filter={"run_id": run_id}, namespace=namespace)
        except Exception as e:
            logger.warning("Delete skipped (namespace may not exist): %s", e, extra=log_context)

        pinecone_vectors = []
        for i in range(len(docs)):
            pinecone_vectors.append({
                "id": f"{run_id}_chunk_{i}",
                "values": vectors[i],
                "metadata": docs[i].metadata,
            })

        index.upsert(vectors=pinecone_vectors, namespace=namespace)

        log_context["namespace"] = namespace
        new_state["brd_embedded"] = True
        logger.info("Safe upsert completed", extra=log_context)
        logger.info("END: _chunk_and_embed", extra=log_context)
        return new_state

    except Exception as e:
        # Best-effort: embedding failures should not block the rest of the pipeline
        # (especially for structured/SFTP runs where we still want KPI extraction).
        logger.info("BRD semantic index deferred; requirement extraction will use parsed BRD content: %s", e, extra=log_context)
        logger.warning(traceback.format_exc(), extra=log_context)
        return _mark_embedding_skipped(new_state, brd_embedded=False)


def _embed_schema_metadata(state: Stage01State) -> Stage01State:
    new_state = state.copy()
    run_id = new_state.get("run_id", "unknown")

    log_context = {
        "run_id": run_id,
        "node": "embed_schema_metadata",
    }

    logger.info("START: _embed_schema_metadata", extra=log_context)

    try:
        if new_state.get("status") == "FAILED":
            return new_state

        model = _get_embedding_model(log_context=log_context)
        if model is None:
            logger.info("Schema semantic index deferred; nomination will use catalog and lexical matching", extra=log_context)
            return _mark_embedding_skipped(new_state, schema_embedded=False, schema_columns_count=0)

        source_databases = new_state.get("source_databases", [])
        if not source_databases:
            logger.info("No source databases selected for schema semantic index", extra=log_context)
            return new_state

        pc = Pinecone(api_key=pinecone_conf.get("api_key") or os.getenv("PINECONE_API_KEY"))
        schema_index_name = os.getenv("PINECONE_SCHEMA_INDEX_NAME") or pinecone_conf.get("schema_index_name") or "metadata"
        index = pc.Index(schema_index_name)

        namespace = "schema"

        all_vectors = []

        source_schema = os.getenv("AZURE_SQL_SOURCE_SCHEMA") or db_conf.get("source_schema") or "dbo"

        for db in source_databases:
            logger.info(f"Fetching schema from DB: {db}", extra=log_context)

            query = """
                SELECT TABLE_SCHEMA, TABLE_NAME, COLUMN_NAME
                FROM INFORMATION_SCHEMA.COLUMNS
                WHERE TABLE_SCHEMA = ?
            """

            rows = execute_source_sql(db, query, (source_schema,))
            if not rows:
                logger.warning(
                    "No columns returned for %s.%s (check AZURE_SQL_SOURCE_* creds/permissions/schema)",
                    db,
                    source_schema,
                    extra=log_context,
                )

            texts = []
            metadata_list = []

            for row in rows:
                schema = row.TABLE_SCHEMA
                table = row.TABLE_NAME
                column = row.COLUMN_NAME

                # ðŸ”¥ semantic-friendly sentence
                text = f"Table {table} contains column {column}"

                texts.append(text)

                metadata_list.append({
                    "database_name": db,
                    "schema_name": schema,
                    "table_name": table,
                    "column_name": column,
                    "type": "schema"
                })

            if not texts:
                continue

            logger.info(f"Embedding {len(texts)} columns from {db}.{source_schema}", extra=log_context)

            vectors = model.embed_documents(texts)

            for i in range(len(vectors)):
                vec_id = f"{db}_{metadata_list[i]['table_name']}_{metadata_list[i]['column_name']}"

                all_vectors.append({
                    "id": vec_id,
                    "values": vectors[i],
                    "metadata": metadata_list[i],
                })

        if not all_vectors:
            logger.warning("No schema vectors generated", extra=log_context)
            return _mark_embedding_skipped(new_state, schema_embedded=False, schema_columns_count=0)

        logger.info(f"Upserting {len(all_vectors)} schema vectors", extra=log_context)

        # Optional: clear old schema embeddings
        try:
            index.delete(delete_all=True, namespace=namespace)
        except Exception:
            pass

        index.upsert(vectors=all_vectors, namespace=namespace)

        new_state["schema_embedded"] = True
        new_state["schema_columns_count"] = len(all_vectors)
        logger.info("END: _embed_schema_metadata", extra=log_context)
        return new_state

    except Exception as e:
        # Best-effort: schema embedding is helpful but not required to keep the run moving.
        logger.info("Schema semantic index deferred; nomination will use catalog and lexical matching: %s", e, extra=log_context)
        logger.warning(traceback.format_exc(), extra=log_context)
        return _mark_embedding_skipped(new_state, schema_embedded=False, schema_columns_count=0)

def ingestion_node(state: Stage01State) -> Stage01State:
    new_state = _copy_state(state)
    source = str(new_state.get("source") or "").lower()

    log_context = {
        "run_id": new_state.get("run_id", "unknown"),
        "node": "ingestion_node",
        "source": source or "database",
    }

    logger.info("START ingestion_node", extra=log_context)

    try:
        new_state = _parse_input(new_state)

        new_state = _acquire_and_validate_brd(new_state)
        if new_state.get("status") == "FAILED":
            logger.warning("Stopped at validation", extra=log_context)
            return new_state

        new_state = _estimate_and_fingerprint(new_state)
        if new_state.get("status") == "FAILED":
            logger.warning("Stopped at fingerprint", extra=log_context)
            return new_state

        new_state = _validate_budget(new_state)
        if new_state.get("status") == "FAILED":
            logger.warning("Stopped at budget check", extra=log_context)
            return new_state

        if source == "sftp":
            new_state["brd_embedded"] = False
            new_state["schema_embedded"] = False
            logger.info("Using file-source catalog analysis path", extra=log_context)
        else:
            new_state = _chunk_and_embed(new_state)
            if new_state.get("status") == "FAILED":
                logger.warning("Stopped at BRD embedding", extra=log_context)
                return new_state

            new_state = _embed_schema_metadata(new_state)
            if new_state.get("status") == "FAILED":
                logger.warning("Stopped at schema embedding", extra=log_context)
                return new_state

        logger.info("END ingestion_node", extra=log_context)
        return new_state

    except Exception as e:
        log_context["error"] = str(e)
        logger.error("ERROR in ingestion_node", extra=log_context)

        new_state.update({
            "status": "FAILED",
            "error": f"Ingestion failed: {str(e)}",
        })
        return new_state
